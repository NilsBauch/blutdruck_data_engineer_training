from docx import Document
import os

template_path = r"c:\Users\nilsb\OneDrive\Nils\weiterbildung\DataEngineer\Projektarbeit\Blutdruck\docs\Template für Speicherplatzberechnung.docx"

if not os.path.exists(template_path):
    print(f"Error: Template not found at {template_path}")
else:
    doc = Document(template_path)
    print("--- PARAGRAPHS ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"Para {i}: {text}")
    
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for r_idx, row in enumerate(table.rows):
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.replace("|", "").strip():
                print(f"Row {r_idx}: {row_text}")
