from docx import Document
import os

template_path = r"c:\Users\nilsb\OneDrive\Nils\weiterbildung\DataEngineer\Projektarbeit\Blutdruck\docs\Template für Speicherplatzberechnung.docx"

if not os.path.exists(template_path):
    print(f"Error: Template not found at {template_path}")
else:
    doc = Document(template_path)
    print("--- DOCUMENT CONTENT ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row in table.rows:
            print(" | ".join(cell.text.strip() for cell in row.cells))
