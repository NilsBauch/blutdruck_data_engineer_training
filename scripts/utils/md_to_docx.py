# ==============================================================================
# SCRIPT: md_to_docx.py
# BESCHREIBUNG: Konvertiert die Design-Dokumentation von Markdown (.md) in das 
#               Word-Format (.docx). Unterstützt Tabellen, Bilder und Listen.
# AUFRUF: py scripts/utils/md_to_docx.py
# VORAUSSETZUNG: Bibliothek 'python-docx'
# ERGEBNIS: Eine .docx Datei im Verzeichnis 'Architektur/Design/'.
# ==============================================================================

import os
from docx import Document
from docx.shared import Inches
import re

# --- KONFIGURATION & PFADE ---
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
MD_FILE = os.path.join(BASE_DIR, 'Architektur', 'Design', 'Design_Dokumentation.md')
DOCX_FILE = os.path.join(BASE_DIR, 'Architektur', 'Design', 'Design_Dokumentation.docx')
IMAGE_DIR = os.path.join(BASE_DIR, 'Architektur', 'images')

def convert_md_to_docx():
    print(f"Lese {MD_FILE}...")
    if not os.path.exists(MD_FILE):
        print("Fehler: MD Datei nicht gefunden.")
        return

    doc = Document()
    
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []
    in_mermaid = False

    for line in lines:
        line = line.strip('\n')
        
        # Mermaid code überspringen (wird im Anhang meist nicht benötigt oder nur als Text)
        if line.startswith('```mermaid'):
            in_mermaid = True
            continue
        if in_mermaid:
            if line.startswith('```'):
                in_mermaid = False
            continue

        # Überschriften
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        
        # Bilder einbetten ![alt](../images/path.png)
        elif line.startswith('!['):
            match = re.search(r'!\[.*?\]\((.*?)\)', line)
            if match:
                img_rel_path = match.group(1)
                img_name = os.path.basename(img_rel_path)
                img_abs_path = os.path.join(IMAGE_DIR, img_name)
                if os.path.exists(img_abs_path):
                    print(f"Bette Bild ein: {img_name}")
                    doc.add_picture(img_abs_path, width=Inches(6.0))
        
        # Tabellen
        elif '|' in line and not line.startswith('---'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Trennlinie überspringen (| :--- |)
            if re.match(r'^[|\s\-:]+$', line):
                continue
                
            cells = [c.strip() for c in line.split('|') if c.strip() or line.split('|').index(c) > 0 and line.split('|').index(c) < len(line.split('|')) - 1]
            if any(cells):
                table_data.append(cells)
        
        # Leere Zeile beendet Tabelle
        elif line.strip() == '' and in_table:
            # Tabelle in Word erstellen
            if len(table_data) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for r, row_cells in enumerate(table_data):
                    for c, cell_text in enumerate(row_cells):
                        try:
                            table.cell(r, c).text = cell_text
                        except:
                            pass
            in_table = False
            table_data = []
            doc.add_paragraph()
        
        # Normaler Text & Listen
        elif line.strip() != '':
            if line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                # Einfache Fett-Ersetzung
                clean_line = line.replace('**', '')
                doc.add_paragraph(clean_line)
        else:
            if not in_table:
                doc.add_paragraph()

    print(f"Speichere {DOCX_FILE}...")
    doc.save(DOCX_FILE)
    print("Erfolg!")

if __name__ == "__main__":
    convert_md_to_docx()
