# ==============================================================================
# SCRIPT: generate_storage_calculation.py
# BESCHREIBUNG: Berechnet den geschätzten Speicherplatzbedarf für Staging und DWH
#               basierend auf einem 1-Jahres-Szenario und erstellt Word-Dateien.
# AUFRUF: py scripts/utils/generate_storage_calculation.py
# VORAUSSETZUNG: Bibliotheken 'python-docx' (Installation: py -m pip install python-docx)
# ERGEBNIS: Zwei ausgefüllte Word-Dokumente im Ordner 'docs/'.
# ==============================================================================

import os
from docx import Document
from docx.shared import Pt, RGBColor

# --- KONFIGURATION & PFADE ---
# BASE_DIR zeigt auf das Hauptverzeichnis (Blutdruck/)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
TEMPLATE_PATH = os.path.join(BASE_DIR, 'docs', 'Template für Speicherplatzberechnung.docx')
OUTPUT_DIR = os.path.join(BASE_DIR, 'docs')

# --- BERECHNUNGS-PARAMETER (User-Vorgaben) ---
DAYS = 365                       # Zeitraum: 1 Jahr
MEASUREMENTS_PER_DAY = 5         # Frequenz: 5 Blutdruck-Messungen pro Tag
SCD_CHANGES_PER_MONTH = 2        # Historisierung: 2 Profil-Änderungen pro Monat
SCD_TOTAL = SCD_CHANGES_PER_MONTH * 12

def calculate_bytes(tables, title, output_filename):
    if not os.path.exists(TEMPLATE_PATH):
        print(f"Error: Template not found at {TEMPLATE_PATH}")
        return

    doc = Document(TEMPLATE_PATH)
    
    # 1. Überschrift und Annahmen hinzufügen
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(f'Speicherplatzberechnung: {title}')
    run_title.bold = True
    run_title.font.size = Pt(24)
    
    p = doc.add_paragraph()
    run = p.add_run("Berechnungsgrundlage (Annahmen):\n")
    run.bold = True
    p.add_run(f"- Benutzer: 1\n")
    p.add_run(f"- Zeitraum: 1 Jahr ({DAYS} Tage)\n")
    p.add_run(f"- Messungen pro Tag: {MEASUREMENTS_PER_DAY}\n")
    p.add_run(f"- Historisierung: {SCD_CHANGES_PER_MONTH} Änderungen/Monat (DWH)")

    # Vorlage hat normalerweise eine Tabelle (Table 0)
    # Wir löschen die Beispiel-Zeilen (außer Header) und füllen sie neu
    if not doc.tables:
        table = doc.add_table(rows=1, cols=15)
        table.style = 'Table Grid'
    else:
        table = doc.tables[0]
        # Beispielzeilen entfernen (wir behalten den Header bei Row 0)
        while len(table.rows) > 1:
            row = table.rows[-1]
            row._element.getparent().remove(row._element)

    total_project_bytes = 0

    for table_name, columns, row_multiplier in tables:
        bytes_per_ds = 0
        anz_ds = row_multiplier
        
        # Sektion für Tabelle
        for col_data in columns:
            name, dtype, length, storage_bytes, pk, fk, idx, not_null = col_data
            bytes_per_ds += storage_bytes
            
            row = table.add_row()
            cells = row.cells
            cells[0].text = table_name
            cells[1].text = name
            cells[2].text = str(anz_ds)
            cells[3].text = str(storage_bytes)
            cells[4].text = dtype
            cells[5].text = str(length) if length else ""
            cells[6].text = "X" if pk else ""
            cells[7].text = "X" if fk else ""
            cells[11].text = "X" if idx or pk else ""
            cells[13].text = "X" if not_null else ""

        # Summenzeilen pro Relation
        # Oelquelle (1 DS)
        r_sum_ds = table.add_row()
        r_sum_ds.cells[0].text = f"{table_name} 1 DS"
        r_sum_ds.cells[3].text = str(bytes_per_ds)
        # Tabelle gesamt
        r_total = table.add_row()
        r_total.cells[0].text = f"{table_name} gesamt"
        table_total = bytes_per_ds * anz_ds
        r_total.cells[3].text = str(table_total)
        
        total_project_bytes += table_total

    # Abschließende Berechnungen
    doc.add_paragraph("\n")
    doc.add_paragraph(f"Gesamtsumme in Bytes: {total_project_bytes}")
    mega_bytes = total_project_bytes / (1024 * 1024)
    doc.add_paragraph(f"Gesamtsumme in MB: {mega_bytes:.4f} MB")

    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"Erfolg: {output_filename} erstellt.")

# Definitionen für Staging (blutdruck_input.db)
# Relation | Feldname | Anz.DS | Storage Bytes | Datentyp | Länge | PK | FK | ID | Def | Comp | Idx | Unq | NN | check
# columns: (name, dtype, length, storage_bytes, pk, fk, idx, not_null)
staging_tables = [
    ("master_medications", [
        ("med_id", "INTEGER", None, 4, True, False, True, True),
        ("name", "TEXT", 100, 100, False, False, False, True),
        ("dose_mg", "REAL", None, 8, False, False, False, True),
        ("description", "TEXT", 200, 200, False, False, False, False),
    ], 20),
    ("master_lifestyle", [
        ("user_id", "INTEGER", None, 4, True, False, True, True),
        ("name", "TEXT", 50, 50, False, False, False, False),
        ("age", "INTEGER", None, 4, False, False, False, False),
        ("gender", "VARCHAR", 1, 1, False, False, False, False),
        ("is_smoker", "BOOLEAN", None, 1, False, False, False, False),
        ("movement_type", "TEXT", 20, 20, False, False, False, False),
        ("raw_data_folder", "TEXT", 100, 100, False, False, False, False),
    ], 1),
    ("user_medication_plan", [
        ("plan_id", "INTEGER", None, 4, True, False, True, True),
        ("user_id", "INTEGER", None, 4, False, True, True, True),
        ("medication_id", "INTEGER", None, 4, False, True, True, True),
        ("time_of_day", "TEXT", 20, 20, False, False, False, False),
        ("is_active", "BOOLEAN", None, 1, False, False, False, False),
    ], 5),
    ("raw_blood_pressure", [
        ("bp_id", "INTEGER", None, 4, True, False, True, True),
        ("user_id", "INTEGER", None, 4, False, True, True, True),
        ("timestamp", "TEXT", 20, 20, False, False, True, True),
        ("systolic", "INTEGER", None, 4, False, False, False, False),
        ("diastolic", "INTEGER", None, 4, False, False, False, False),
        ("pulse", "INTEGER", None, 4, False, False, False, False),
        ("is_manual", "BOOLEAN", None, 1, False, False, False, False),
    ], DAYS * MEASUREMENTS_PER_DAY),
    ("raw_activity_daily", [
        ("activity_id", "INTEGER", None, 4, True, False, True, True),
        ("user_id", "INTEGER", None, 4, False, True, True, True),
        ("date", "TEXT", 10, 10, False, False, True, True),
        ("steps", "INTEGER", None, 4, False, False, False, False),
        ("activity_minutes", "INTEGER", None, 4, False, False, False, False),
        ("weight_kg", "REAL", None, 8, False, False, False, False),
    ], DAYS),
]

# Definitionen für DWH (blutdruck_dwh.db)
dwh_tables = [
    ("dim_user", [
        ("user_id", "INTEGER", None, 4, True, False, True, True),
        ("gender", "VARCHAR", 1, 1, False, False, False, False),
        ("age", "INTEGER", None, 4, False, False, False, False),
    ], 1),
    ("dim_medication", [
        ("med_key", "INTEGER", None, 4, True, False, True, True),
        ("med_id", "INTEGER", None, 4, False, False, True, False),
        ("name", "TEXT", 100, 100, False, False, False, False),
        ("dosage_mg", "REAL", None, 8, False, False, False, False),
        ("category", "TEXT", 50, 50, False, False, False, False),
        ("SCD_valid_from", "DATE", 10, 10, False, False, False, False),
        ("SCD_valid_to", "DATE", 10, 10, False, False, False, False),
    ], 1 + SCD_TOTAL),
    ("dim_lifestyle", [
        ("lifestyle_key", "INTEGER", None, 4, True, False, True, True),
        ("user_id", "INTEGER", None, 4, False, False, True, False),
        ("is_smoker", "BOOLEAN", None, 1, False, False, False, False),
        ("movement_type", "TEXT", 20, 20, False, False, False, False),
        ("SCD_valid_from", "DATE", 10, 10, False, False, False, False),
        ("SCD_valid_to", "DATE", 10, 10, False, False, False, False),
    ], 1 + SCD_TOTAL),
    ("dim_date", [
        ("date_key", "INTEGER", None, 4, True, False, True, True),
        ("full_date", "DATE", 10, 10, False, False, False, False),
        ("day", "INTEGER", None, 4, False, False, False, False),
        ("month", "INTEGER", None, 4, False, False, False, False),
        ("year", "INTEGER", None, 4, False, False, False, False),
        ("day_name", "TEXT", 20, 20, False, False, False, False),
        ("is_weekend", "BOOLEAN", None, 1, False, False, False, False),
    ], DAYS),
    ("fact_health_metrics", [
        ("fact_id", "INTEGER", None, 4, True, False, True, True),
        ("user_id", "INTEGER", None, 4, False, False, True, False),
        ("date_key", "INTEGER", None, 4, False, True, True, False),
        ("time_key", "TEXT", 5, 5, False, False, True, False),
        ("med_key", "INTEGER", None, 4, False, True, True, False),
        ("lifestyle_key", "INTEGER", None, 4, False, True, True, False),
        ("systolic", "INTEGER", None, 4, False, False, False, False),
        ("diastolic", "INTEGER", None, 4, False, False, False, False),
        ("pulse", "INTEGER", None, 4, False, False, False, False),
        ("steps_hourly", "INTEGER", None, 4, False, False, False, False),
        ("weight_kg", "REAL", None, 8, False, False, False, False),
        ("activity_minutes", "INTEGER", None, 4, False, False, False, False),
        ("is_post_medication", "BOOLEAN", None, 1, False, False, False, False),
        ("pulse_pressure", "INTEGER", None, 4, False, False, False, False),
        ("load_timestamp", "DATETIME", 20, 20, False, False, False, False),
    ], DAYS * MEASUREMENTS_PER_DAY),
]

if __name__ == "__main__":
    calculate_bytes(staging_tables, "Staging-Bereich (Ingestion)", "Speicherplatzberechnung_Staging.docx")
    calculate_bytes(dwh_tables, "Data Warehouse (Analytics)", "Speicherplatzberechnung_DWH.docx")
